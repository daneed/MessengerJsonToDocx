import os, json, datetime, re, textwrap
import argparse, pathlib
from docx import Document
from docx.shared import RGBColor, Cm, Pt
from PIL import Image
import dominate
from dominate.tags import *
from abc import ABC, abstractmethod
import shutil



class ProcessorBase (ABC):
    def __init__(self, file):
        self.file=file
        self.senderNameToColorDict = dict()
        self.bankAccountRegexObject = re.compile(r'.*([1-9][0-9]{7})([\-|\s]*)([0-9]{8})([\-|\s]*)([0-9]{8}).*', re.IGNORECASE|re.MULTILINE)
        self.phoneNumberRegexObject = re.compile(r'.*(\+36|06)([\s|\-]*)([0-9]{2})([\s|\-]*)([0-9]{3})([\s|\-]*)([0-9]{4}).*', re.IGNORECASE|re.MULTILINE)
        self.color = None

    def SetMainColor(self, color):
        self.color = color

    @abstractmethod
    def AddTitle(self):
        pass

    @abstractmethod
    def InitOneConversationData(self, color):
        pass

    @abstractmethod
    def AddDate(self, date):
        pass

    @abstractmethod
    def InitNameAndDataRow(self):
        pass

    @abstractmethod
    def AddName(self, name):
        pass

    @abstractmethod
    def AddMessageText(self, text):
        pass

    @abstractmethod
    def AddPicture(self, picturePath, width, height):
        pass

    @abstractmethod
    def AddMediaNameWithDetails(self, mediaName, details):
        pass

    @abstractmethod
    def AddUrl(self, url):
        pass

    @abstractmethod
    def Save(self):
        pass

    def Do(self, prefix, printCallback):
        with open(str(self.file), 'r', encoding="utf-8") as j:
            self.jsonContent = json.loads(j.read())

        if not 'messages' in self.jsonContent or len (self.jsonContent['messages']) == 0:
            print (f'{prefix} File does not contain any messages.')
            return
        
        self.AddTitle()

        for participant in self.jsonContent['participants']:
            self._SenderNameToColor(prefix, participant)

        forbiddenStringRegexpPatterns = self._GetForbiddenStringRegexpPatterns(prefix)
        replacements = dict()

        count = 0
        for message in self.jsonContent['messages']:
            printCallback(count, False)
            count += 1
            type = message['type']
            senderName = message['senderName']
            
            self.InitOneConversationData()
            myDate = datetime.datetime.fromtimestamp(int (message['timestamp']/1000.0))
            self.AddDate(myDate)

            color = self._SenderNameToColor(prefix, senderName)

            self.SetMainColor(color)

            self.InitNameAndDataRow()

            self.AddName(senderName)

            if 'text' in message and type != 'media':
                messageText = str (message['text']).strip()
                prettyMessageText = messageText.replace(f'\n', f'\n{prefix}')

                for pattern in forbiddenStringRegexpPatterns:
                    match = pattern.search (messageText)
                    if match:
                        if pattern in replacements:
                            replacement = replacements[pattern]
                        else:
                            print (f"\n{prefix}NOTE: FOUND\n{prefix} '{match.group()}' \n{prefix}IN\n{prefix}'{prettyMessageText}'")
                            replacement = input(f'Enter replacement for {match.group()} : ').strip()
                            applyForAllLater = input ('Enter y for apply all later existance: ').upper().strip()
                            if applyForAllLater:
                                replacements[pattern] = replacement

                        if len (replacement) > 0:
                            messageText = re.sub (pattern, replacement, messageText)

                splittedWordsWithWhitespaces = list ()
                words_wo_sp = messageText.split (" ")
                for i in range(len(words_wo_sp)):
                    if i < len(words_wo_sp) - 1: words_wo_sp[i] += " "
                    words_wo_n = words_wo_sp[i].split ("\n")
                    for j in range (len (words_wo_n)):
                         if j < len(words_wo_n) - 1: words_wo_n[j] += "\n"
                         words_wo_r = words_wo_n[j].split("\r")
                         for k in range (len (words_wo_r)):
                              if k < len(words_wo_r) - 1: words_wo_r[k] += "\r"
                              words_wo_t = words_wo_r[k].split("\r")
                              for l in range (len (words_wo_t)):
                                  if l < len(words_wo_t) - 1: words_wo_t[l] += "\t"
                                  splittedWordsWithWhitespaces.append (words_wo_t[l])

                tmpWords = list()
                for i in range (len(splittedWordsWithWhitespaces)):
                    tmp = splittedWordsWithWhitespaces[i].split ("-")
                    for j in range (len(tmp)):
                        currWord = tmp[j]
                        if len (currWord) > 80:
                            currWord = " ".join (textwrap.wrap (currWord, width=80))
                        if j < len (tmp) -1:
                            tmpWords.append (currWord + "-")
                        else:
                            tmpWords.append(currWord)
                messageText = ''.join (tmpWords)
                self.AddMessageText(messageText)

            if type == 'media':
                for media in message['media']:
                    if 'uri' in media:
                        mediaName = media['uri'][2:]
                        picturePath = self.file.absolute().parent/mediaName
                        if (picturePath.suffix in ['.jpeg', '.gif']):
                            im = Image.open(str (picturePath))
                            imWidth, imHeight = im.size
                            self.AddPicture(str(picturePath), imWidth, imHeight)
                        elif picturePath.suffix:
                            details = input(f"\nPlease enter details of {mediaName}. You can leave empty: ")
                            self.AddMediaNameWithDetails(mediaName, details)
                        else:
                            self.AddUrl(media['uri'])

        self.Save()
        printCallback(count, True)
    
    def _SenderNameToColor(self, prefix, senderName):

        if senderName not in self.senderNameToColorDict:
            color=input(f"{prefix}Enter color for {senderName}. #RRGGBB, or R for red, G for green, B for blue: ")
            color = color.upper()
            if color == 'R':
                self.senderNameToColorDict[senderName] = RGBColor(0xFF, 0x00, 0x00)
            elif color == 'G':
                self.senderNameToColorDict[senderName] = RGBColor(0x00, 0xFF, 0x00)
            elif color == 'B':
                self.senderNameToColorDict[senderName] = RGBColor(0x00, 0x00, 0xFF)
            else:
                if not color.startswith('#'):
                    raise Exception ("Illegal color format!")
                else:
                    try:
                        hexR = int (f"0x{color[1:3]}", 16)
                        hexG = int (f"0x{color[3:5]}", 16)
                        hexB = int (f"0x{color[5:]}", 16)
                        self.senderNameToColorDict[senderName] = RGBColor(hexR, hexG, hexB)
                    except:
                        raise Exception ("Illegal color format!")
            print (f"{prefix}This color will be used for {senderName}: #{str (self.senderNameToColorDict[senderName])}")
        return self.senderNameToColorDict[senderName]
    
    def _GetForbiddenStringRegexpPatterns (self, prefix):
        canContinue = True
        forbiddenStrings = list()
        while (canContinue):
            forbiddenString = input (f'{prefix}Enter a forbidden string, or press enter to continue: ').strip()
            if len(forbiddenString) > 0:
                forbiddenStrings.append(re.compile (rf'{forbiddenString}', re.IGNORECASE|re.MULTILINE))
            else:
                canContinue=False

        doMask=input (f'{prefix}Max bank accounts? press Y to yes: ').strip().upper()
        if doMask == 'Y':
            forbiddenStrings.append (self.bankAccountRegexObject)

        doMask=input (f'{prefix}Max phone numbers? press Y to yes: ').strip().upper()
        if doMask == 'Y':
            forbiddenStrings.append (self.phoneNumberRegexObject)
        
        return forbiddenStrings


class ProcessorWithDocXOutput (ProcessorBase):
    def __init__(self, file):
        super().__init__(file)
        self.document = Document()

        section = self.document.sections[0]
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

        style = self.document.styles['Normal']
        style.font.name = 'Consolas'
        style.font.size = Pt(10)

    def AddTitle(self):
        self.document.add_heading(f"Messenger chat közöttük: {", ".join (self.jsonContent['participants'])}")

    def InitOneConversationData(self):
        para0 = self.document.add_paragraph('')
        para0.style.paragraph_format.space_before= Pt(0)
        para0.style.paragraph_format.space_after= Pt(0)
        para0.style.next_paragraph_style = None

    def AddDate(self, date):
        self.document.add_paragraph(f'[{str(date)}]')

    def InitNameAndDataRow(self):
        section = self.document.sections[0]
        table = self.document.add_table(rows=1, cols=2)
        table.autofit = True
        self.nameCell = table.cell(0, 0)
        self.dataCell = table.cell(0, 1)
        self.nameCell.width = Cm(3.5)
        self.dataCell.width = section.page_width - self.nameCell.width - section.right_margin - section.left_margin

    def AddName(self, name):
        senderNameRun = self.nameCell.paragraphs[0].add_run(f'{name}:')
        senderNameRun.font.bold = True
        senderNameRun.font.color.rgb = self.color

    def AddMessageText(self, text):
        self.dataCell.paragraphs[0].add_run(f'{text}')

    def AddPicture(self, picturePath, imWidth, imHeight):
        width = None
        height = None
        if imWidth > imHeight:
            width = Cm(10)
        elif imHeight > 500:
            height = Cm(5)
        mediaRun = self.dataCell.paragraphs[0].add_run()
        mediaRun.font.italic = True
        mediaRun.add_picture(str (picturePath), width, height)

    def AddMediaNameWithDetails(self, mediaName, details):
        if details:
            mediaRun = self.dataCell.paragraphs[0].add_run()
            mediaRun.font.italic = True
            mediaRun.add_text(f"<{mediaName}>: {details}")
        else:
            mediaRun = self.dataCell.paragraphs[0].add_run()
            mediaRun.font.italic = True
            mediaRun.add_text(f"<{mediaName}>.")

    def AddUrl(self, url):
        mediaRun = self.dataCell.paragraphs[0].add_run()
        mediaRun.font.italic = True
        mediaRun.add_text(f"<{url}>")


    def Save(self):
        docName = (pathlib.Path (os.getcwd()) / self.file.name).with_suffix('.docx')
        self.document.save(docName)


class ProcessorWithHtmlOutput(ProcessorBase):
    def __init__(self, file):
        super().__init__(file)
        self.document = dominate.document()
        
        with self.document.head:
            link(rel='stylesheet', href='messengerStyle.css')
            meta(charset="UTF-8")

        self.container = self.document.add(div(cls="container"))

    def AddTitle(self):
        self.document.title = f"Messenger chat közöttük: {", ".join (self.jsonContent['participants'])}"

    def InitOneConversationData(self):
        self.oneConversation = self.container.add(div(cls="oneConversation"))

    def AddDate(self, date):
        self.oneConversation.add(div(f"[{str(date)}]", cls="date"))

    def InitNameAndDataRow(self):
        self.oneConversationData = self.oneConversation.add(div(cls="oneConversationData"))
        self.nameCell = self.oneConversationData.add(div(cls="nameCell"))
        self.theColor = '#{:02X}{:02X}{:02X}'.format(self.color[0],self.color[1],self.color[2])
        self.dataCellWrapper = self.oneConversationData.add(div(cls="dataCellWrapper"))

    def AddName(self, name):
        self.nameCell.add(b(name))

    def AddMessageText(self, text):
        dataCell = self.dataCellWrapper.add(div(cls="dataCell",style=f"background-color:{self.theColor};"))
        dataCell.add(text)

    def AddPicture(self, picturePath, imWidth, imHeight):
        imageFolderName = pathlib.Path(picturePath).parent.name
        pictureName = pathlib.Path(picturePath).name
        imageFolderPath = pathlib.Path (os.getcwd()) / imageFolderName
        imageFolderPath.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(picturePath, imageFolderPath / pictureName)
        imageCell = self.dataCellWrapper.add(div(cls="imgCell",style=f"background-color:{self.theColor};"))
        imageCell.add(img(src=f"{imageFolderName}/{pictureName}"))

    def AddMediaNameWithDetails(self, mediaName, details):
        dataCell = self.dataCellWrapper.add(div(cls="dataCell",style=f"background-color:{self.theColor};"))
        if details:
            dataCell.add(i(f"<{mediaName}>: {details}"))
        else:
            dataCell.add(i(f"<{mediaName}>"))

    def AddUrl(self, url):
        dataCell = self.dataCellWrapper.add(div(cls="dataCell",style=f"background-color:{self.theColor};"))
        dataCell.add(a(f"<{url}>", href= f"{url}"))

    def Save(self):
        docName = (pathlib.Path (os.getcwd()) / self.file.name).with_suffix('.html')
        with open(str(docName), 'w', encoding="utf-8") as h:
            h.write(str(self.document))

if __name__ == "__main__":
    parser = argparse.ArgumentParser ("Messenger Json to Docx converter")
    parser.add_argument("-f", "--fileOrFolder", help="The input json file or the container folder", metavar="JSONFILE", required=True)
    parser.add_argument("--html", help="Set this flag, to produce html output, instead of docs!", action="store_true")
    args = parser.parse_args()
    path = pathlib.Path(args.fileOrFolder)

    def printFunction (prefix, count, isEnd):
        if isEnd:
            print ("DONE", flush=True)
        elif count == 0:
            print (f"{prefix}Processing messages...", end="", flush=True)
        elif count % 100 == 0:
            print (".", end="", flush=True)

    if path.is_file() and path.suffix == ".json":
        print (f"Processing file: {path.name}...", flush=True)
        processor = ProcessorWithDocXOutput(path) if not args.html else ProcessorWithHtmlOutput(path)
        processor.Do ("   ", lambda count, isEnd: printFunction ("   ", count, isEnd))
        print (f"Processing file: {path.name}...DONE", flush=True)

    elif path.is_dir ():
        print (f"Processing directory: {path.name}...")
        for subPath in path.iterdir():
            if subPath.is_file() and subPath.suffix == ".json":
                print (f"   Processing file: {subPath.name}...", flush=True)
                processor = ProcessorWithDocXOutput(subPath) if not args.html else ProcessorWithHtmlOutput(subPath)
                processor.Do (
                    "      ",
                    lambda count, isEnd: printFunction ("   ", count, isEnd))
                print (f"   Processing file: {subPath.name}...DONE", flush=True)
        print (f"Processing directory {path.name}...DONE")
    else:
        raise Exception(f"File {path} is not a json file!")
